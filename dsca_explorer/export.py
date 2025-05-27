"""
================================================================================
DSCA Explorer Export Module - Change Log
================================================================================

BEFORE:
-------
- The module only supported exporting raw layer data (via export_layers).
- Supported formats: csv, xlsx, json, txt, docx, pdf.
- No support for exporting change logs or change records.
- No standardized export for field-level change detection or audit trails.
- No integration with ChangeRecord dataclass or cache change detection.

AFTER:
------
- Added export_changes function to export a list of ChangeRecord objects
  (from cache.py) in multiple formats (csv, xlsx, json, txt, docx, pdf).
- Standardized change log export columns: source, layer_id, change_type,
  changed_fields, detection_time.
- export_changes expects ChangeRecord objects with a .to_serializable() method.
- Both export_changes (for change logs) and export_layers (for raw layers)
  are now available in the module.
- All necessary imports and dependencies included.
- Ready for CLI and test integration.

================================================================================
"""


import pandas as pd
import json
from pathlib import Path
from typing import List, Any
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

import requests
import urllib3
import re
from html import unescape
from urllib.parse import urlparse

# Suppress only if verify=False is used
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# --- Change Log Export ---

def export_changes(change_records: List[Any], format: str, output_path: Path):
    """
    Export a list of ChangeRecord objects to the specified format.
    Supported formats: csv, json, xlsx, txt, docx, pdf
    """
    # Convert ChangeRecord objects to dicts (assume .to_serializable() method)
    rows = [c.to_serializable() for c in change_records]

    columns = ["source", "layer_id", "change_type", "changed_fields", "detection_time"]

    if format == "csv":
        pd.DataFrame(rows)[columns].to_csv(output_path, index=False)
    elif format == "xlsx":
        pd.DataFrame(rows)[columns].to_excel(output_path, index=False)
    elif format == "json":
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(rows, f, indent=2, ensure_ascii=False)
    elif format == "txt":
        with open(output_path, "w", encoding="utf-8") as f:
            for row in rows:
                for col in columns:
                    f.write(f"{col}: {row.get(col, '')}\n")
                f.write("\n")
    elif format == "docx":
        doc = Document()
        doc.add_heading("DSCA Change Log Export", 0)
        for row in rows:
            doc.add_heading(row.get('layer_id', ''), level=1)
            for col in columns:
                doc.add_paragraph(f"{col}: {row.get(col, '')}")
        doc.save(str(output_path))
    elif format == "pdf":
        c = canvas.Canvas(str(output_path), pagesize=letter)
        width, height = letter
        y = height - 40
        c.setFont("Helvetica-Bold", 16)
        c.drawString(40, y, "DSCA Change Log Export")
        y -= 30
        c.setFont("Helvetica", 10)
        for row in rows:
            if y < 100:
                c.showPage()
                y = height - 40
            c.setFont("Helvetica-Bold", 12)
            c.drawString(40, y, row.get('layer_id', ''))
            y -= 20
            c.setFont("Helvetica", 10)
            for col in columns:
                if col == "layer_id":
                    continue
                if y < 60:
                    c.showPage()
                    y = height - 40
                value_str = str(row.get(col, ''))
                for line in value_str.splitlines():
                    while len(line) > 100:
                        c.drawString(60, y, line[:100])
                        line = line[100:]
                        y -= 12
                    c.drawString(60, y, line)
                    y -= 12
            y -= 30
        c.save()
    else:
        raise ValueError(f"Unsupported export format: {format}")

# --- Layer Export (unchanged) ---

def clean_html(text):
    """Remove HTML tags and decode entities from text."""
    if not text:
        return ""
    text = re.sub(r'<[^>]+>', '', text)
    return unescape(text).strip()

def get_domain(url):
    try:
        return urlparse(url).netloc
    except Exception:
        return ""

def get_arcgis_compat(endpoint):
    if not endpoint:
        return ""
    if "/MapServer" in endpoint or "/FeatureServer" in endpoint:
        return "ArcGIS Service"
    elif "arcgis.com" in endpoint:
        return "ArcGIS REST API"
    else:
        return "REST API"

def ensure_url_scheme(endpoint):
    """Ensure the endpoint starts with http(s)://"""
    if not endpoint:
        return ""
    if endpoint.startswith("http://") or endpoint.startswith("https://"):
        return endpoint
    return "https://" + endpoint.lstrip("/")

def robust_get(url, params=None, timeout=15):
    """Try to fetch with SSL verification, fallback to verify=False if needed."""
    try:
        return requests.get(url, params=params, timeout=timeout)
    except requests.exceptions.SSLError:
        print(f"SSL error for {url}, retrying without verification (not secure)...")
        return requests.get(url, params=params, timeout=timeout, verify=False)
    except Exception as e:
        print(f"Request error for {url}: {e}")
        return None

def query_feature_attributes(endpoint):
    """Query an ArcGIS endpoint to get actual feature attributes."""
    endpoint = ensure_url_scheme(endpoint)
    if not endpoint:
        return {}

    # Attempt to extract layer_id and service root
    layer_id = None
    service_root = endpoint
    match = re.search(r'(.*(?:MapServer|FeatureServer))/(\d+)', endpoint)
    if match:
        service_root, layer_id = match.group(1), match.group(2)
    else:
        # If not found, try to use as is
        layer_id = "0"

    query_url = f"{service_root}/{layer_id}/query"
    params = {
        "where": "1=1",
        "outFields": "*",
        "returnGeometry": "false",
        "resultRecordCount": 1,
        "f": "json"
    }
    try:
        resp = robust_get(query_url, params=params)
        if resp and resp.status_code == 200:
            data = resp.json()
            if "features" in data and data["features"]:
                attrs = data["features"][0].get("attributes", {})
                if attrs:
                    return attrs
            # Fallback to fields if no features
            if "fields" in data:
                return {f["name"]: f.get("alias", f["name"]) for f in data["fields"]}
        # Fallback: try to get fields from layer metadata
        layer_url = f"{service_root}/{layer_id}"
        resp = robust_get(layer_url, params={"f": "json"})
        if resp and resp.status_code == 200:
            data = resp.json()
            if "fields" in data:
                return {f["name"]: f.get("alias", f["name"]) for f in data["fields"]}
        return {}
    except Exception as e:
        print(f"Error querying features: {e}")
        return {}

def get_property_list(properties):
    """Get a list of property names from a dict."""
    if isinstance(properties, dict):
        return ", ".join(sorted(properties.keys()))
    return "No attributes found"

def export_layers(layers, file_path):
    ext = file_path.split('.')[-1].lower()
    processed_layers = []
    for l in layers:
        endpoint = l.get("endpoint", "")
        # Query feature attributes (with fallback)
        properties = query_feature_attributes(endpoint) if endpoint else l.get("properties", {})
        processed = {
            "Source": l.get("source", ""),
            "Series": l.get("series", ""),
            "Layer Name": l.get("name", ""),
            "Type": l.get("type", ""),
            "Endpoint": endpoint,
            "Domain": get_domain(endpoint),
            "ArcGIS Compatible": get_arcgis_compat(endpoint),
            "Download URL": l.get("download_url", endpoint),
            "Documentation": l.get("documentation", ""),
            "Description": clean_html(l.get("description", "")),
            "Property List": get_property_list(properties),
            "Example Properties": json.dumps(properties, indent=2, ensure_ascii=False)
        }
        processed_layers.append(processed)

    if ext == "csv":
        pd.DataFrame(processed_layers).to_csv(file_path, index=False)
    elif ext == "xlsx":
        pd.DataFrame(processed_layers).to_excel(file_path, index=False)
    elif ext == "json":
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(processed_layers, f, indent=2, ensure_ascii=False)
    elif ext == "txt":
        with open(file_path, "w", encoding="utf-8") as f:
            for l in processed_layers:
                for k, v in l.items():
                    f.write(f"{k}: {v}\n")
                f.write("\n")
    elif ext == "docx":
        doc = Document()
        doc.add_heading("DSCA Layers Export", 0)
        for l in processed_layers:
            doc.add_heading(l['Layer Name'], level=1)
            for k, v in l.items():
                doc.add_paragraph(f"{k}: {v}")
        doc.save(file_path)
    elif ext == "pdf":
        c = canvas.Canvas(file_path, pagesize=letter)
        width, height = letter
        y = height - 40
        c.setFont("Helvetica-Bold", 16)
        c.drawString(40, y, "DSCA Layers Export")
        y -= 30
        c.setFont("Helvetica", 10)
        for l in processed_layers:
            if y < 100:
                c.showPage()
                y = height - 40
            c.setFont("Helvetica-Bold", 12)
            c.drawString(40, y, l['Layer Name'])
            y -= 20
            c.setFont("Helvetica", 10)
            for key, value in l.items():
                if key == "Layer Name":
                    continue
                if y < 60:
                    c.showPage()
                    y = height - 40
                # Wrap long lines
                value_str = str(value)
                for line in value_str.splitlines():
                    while len(line) > 100:
                        c.drawString(60, y, line[:100])
                        line = line[100:]
                        y -= 12
                    c.drawString(60, y, line)
                    y -= 12
            y -= 30
        c.save()
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
